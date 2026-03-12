from fileinput import filename
import streamlit as st  
import requests  
import json 
import os
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import plotly.graph_objects as go
from datetime import datetime

# configure Streamlit page layout and appearance
st.set_page_config(
    page_title="Legal Document Assistant",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

load_dotenv()

API_URL = "https://models.inference.ai.azure.com/chat/completions"
API_KEY = os.getenv("GITHUB_TOKEN")

# allow configuration via Streamlit secrets as well (useful on streamlit.io)
try:
    if "GITHUB_TOKEN" in st.secrets:
        API_KEY = st.secrets["GITHUB_TOKEN"]
except Exception:
    # st.secrets is not available when running locally without a secrets file
    pass

MODEL_NAME = "gpt-4o-mini"

# guard against sending an enormous document (413 errors).
# the endpoint may reject payloads larger than a few kilobytes; 8k chars
# is a conservative per‑request cap.  The code chunks longer text.
MAX_DOC_CHARS = 8000
  

def analyze_document(document_text):
    # handle empty input immediately
    if not document_text:
        return {"error": "Empty document"}

    # ensure API key is available before making any API calls
    if not API_KEY:
        return {
            "error": "API key not configured. Set GITHUB_TOKEN in your local .env file "
                     "or in Streamlit secrets, then rerun the app."
        }

    # if the document exceeds the per-request limit, process it in chunks and
    # merge the partial analyses. recursion ensures each chunk is <= limit.
    if len(document_text) > MAX_DOC_CHARS:
        combined = {"summary": [], "clauses": [], "obligations": [], "risks": [], "keyTerms": []}
        for i in range(0, len(document_text), MAX_DOC_CHARS):
            chunk = document_text[i : i + MAX_DOC_CHARS]
            part = analyze_document(chunk)  # recursive call on shorter text
            if "error" in part:
                return part
            combined["summary"].append(part.get("summary", ""))
            combined["clauses"].extend(part.get("clauses", []))
            combined["obligations"].extend(part.get("obligations", []))
            combined["risks"].extend(part.get("risks", []))
            combined["keyTerms"].extend(part.get("keyTerms", []))
        combined["summary"] = " ".join(s for s in combined["summary"] if s)
        return combined

    prompt = (f"""You are a helpful legal assistant. Please analyze this legal document and 
provide:
1. A simple summary (2-3 sentences)
2. Key clauses explained in simple language (3-5 important ones)
3. What the person must do (obligations)
4. Potential risks or problems
5. Important legal terms with simple definitions
Please respond in JSON format with these keys:- summary: string- clauses: list of objects with "title" and "explanation"- obligations: list of strings- risks: list of strings- keyTerms: list of objects with "term" and "definition"
Here is the document:
{document_text}""")

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }
    data = {
        "model": MODEL_NAME, 
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 2000,
                "temperature": 0.3,  
                "response_format": {"type": "json_object"}
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data)
        if response.status_code == 413:
            # the chunk we sent was still too big -- split again
            half = len(document_text) // 2
            part1 = analyze_document(document_text[:half])
            part2 = analyze_document(document_text[half:])
            if "error" in part1:
                return part1
            if "error" in part2:
                return part2
            merged = {"summary": [], "clauses": [], "obligations": [], "risks": [], "keyTerms": []}
            for p in (part1, part2):
                merged["summary"].append(p.get("summary", ""))
                merged["clauses"].extend(p.get("clauses", []))
                merged["obligations"].extend(p.get("obligations", []))
                merged["risks"].extend(p.get("risks", []))
                merged["keyTerms"].extend(p.get("keyTerms", []))
            merged["summary"] = " ".join(s for s in merged["summary"] if s)
            return merged
        if response.status_code == 200:
            result = response.json()

            ai_text = result["choices"][0]["message"]["content"]

            try:
                analysis = json.loads(ai_text)
                return analysis
            except:
                return {
                    "error": "Model did not retirn JSON",
                    "raw_response": ai_text
                }
        else:
            if response.status_code == 401:
                return {
                    "error": "Authentication failed with status code 401. "
                             "Please verify that your GITHUB_TOKEN is correct and has access "
                             "to GitHub Models / Azure Inference, then restart the app.",
                    "details": response.text,
                }
            return {
                "error": (f"API request failed with status code {response.status_code}"),
                "details": response.text,
            }
    except Exception as e:
        return {
            "error": f"An error occurred: {str(e)}"
        }

def compute_risk_score(analysis):
    """Simple risk meter based on number of identified risks.
    Returns an integer percentage 0-100."""
    risks = analysis.get("risks") or []
    if not risks:
        return 0
    # each risk counts for 20 points up to 100
    score = min(len(risks) * 20, 100)
    return score


def extract_dates_with_context(text):
    """Extracts various date formats with surrounding context snippets."""
    import re

    patterns = [
        r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|'
        r'September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|'
        r'September|October|November|December)\s+\d{4}\b',
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{4}-\d{2}-\d{2}\b',
    ]

    results = []
    seen = set()
    for pat in patterns:
        for match in re.finditer(pat, text, re.IGNORECASE):
            date_str = match.group(0).strip()
            if date_str.lower() in seen:
                continue
            seen.add(date_str.lower())
            start = max(0, match.start() - 80)
            end = min(len(text), match.end() + 80)
            snippet = text[start:end].replace("\n", " ").strip()
            results.append({"date": date_str, "context": snippet})
    return results


def render_deadline_tracker(dates):
    """Render extracted dates as timeline cards in Streamlit."""
    if not dates:
        st.info("📅 No specific dates or deadlines found in this document.")
        return
    st.success(f"**🎯 {len(dates)} date(s) detected** — Review them carefully so you never miss a deadline!")
    for item in dates:
        context = item["context"].replace(item["date"], f"**{item['date']}**")
        st.markdown(f"""
        <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    padding: 1rem; border-radius: 10px; margin: 0.5rem 0; 
                    border-left: 5px solid #ffd700; box-shadow: 0 4px 6px rgba(0,0,0,0.1);'>
            <p style='color: white; margin: 0;'>📌 {context}</p>
        </div>
        """, unsafe_allow_html=True)

def ask_question(document_text, question, conversation_history):
    # ensure we don't blow up the request size; truncate if necessary
    if len(document_text) > MAX_DOC_CHARS:
        document_text = document_text[:MAX_DOC_CHARS]
        # warning to UI is handled by caller if desired

    # ensure API key is available before making any API calls
    if not API_KEY:
        return (
            "API key not configured. Set GITHUB_TOKEN in your local .env file or in "
            "Streamlit secrets, then rerun the app."
        )

    messages = []

    for msg in conversation_history:
        messages.append({
            "role": msg["role"],
            "content": msg["content"]
        })

    current_question = f"""Context: Here is the legal document:
{document_text}
User's question: {question}
Please provide a clear, simple answer that helps them understand the document better."""
    messages.append({
        "role": "user",
        "content": current_question
    })

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }
    data = {
        "model": MODEL_NAME,
        "max_tokens": 1000,
        "messages": messages
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            answer = result["choices"][0]["message"]["content"]
            return answer
        elif response.status_code == 401:
            return (
                "Authentication failed with status code 401. "
                "Please check that your GITHUB_TOKEN is valid and has access to the API."
            )
        else:
            return (
                f"Sorry, I couldn't process your question. "
                f"Error code: {response.status_code}"
            )
    except Exception as e:
        return f"An error occurred: {str(e)}"

# --- additional helpers copied from the full feature set ---
# crime records
CRIME_KEYWORDS = {
    "murder":           ("CRITICAL", "Violent Crime"),
    "homicide":         ("CRITICAL", "Violent Crime"),
    "manslaughter":     ("CRITICAL", "Violent Crime"),
    "assault":          ("HIGH",     "Violent Crime"),
    "battery":          ("HIGH",     "Violent Crime"),
    "kidnapping":       ("CRITICAL", "Abduction"),
    "abduction":        ("CRITICAL", "Abduction"),
    "trafficking":      ("CRITICAL", "Abduction"),
    "hit and run":      ("HIGH",     "Traffic Crime"),
    "hit-and-run":      ("HIGH",     "Traffic Crime"),
    "reckless driving": ("MEDIUM",   "Traffic Crime"),
    "smuggling":        ("CRITICAL", "Organized Crime"),
    "drug trafficking": ("CRITICAL", "Organized Crime"),
    "extortion":        ("HIGH",     "Organized Crime"),
    "narcotics":        ("HIGH",     "Organized Crime"),
    "bribery":          ("HIGH",     "White Collar"),
    "corruption":       ("HIGH",     "White Collar"),
    "fraud":            ("HIGH",     "White Collar"),
    "embezzlement":     ("HIGH",     "White Collar"),
    "money laundering": ("CRITICAL", "White Collar"),
    "robbery":          ("HIGH",     "Property Crime"),
    "burglary":         ("MEDIUM",   "Property Crime"),
    "arson":            ("HIGH",     "Property Crime"),
    "theft":            ("MEDIUM",   "Property Crime"),
    "terrorism":        ("CRITICAL", "National Security"),
    "rape":             ("CRITICAL", "Sexual Offence"),
    "sexual assault":   ("CRITICAL", "Sexual Offence"),
    "stalking":         ("HIGH",     "Sexual Offence"),
    "criminal":         ("MEDIUM",   "General"),
    "felony":           ("HIGH",     "General"),
    "conspiracy":       ("HIGH",     "General"),
    "hacking":          ("MEDIUM",   "Cybercrime"),
    "cyberattack":      ("HIGH",     "Cybercrime"),
}

SEV_STYLE = {
    "CRITICAL": {"color": "#c53030", "bg": "#fff5f5", "icon": "🔴"},
    "HIGH":     {"color": "#c05621", "bg": "#fffaf0", "icon": "🟠"},
    "MEDIUM":   {"color": "#b7791f", "bg": "#fefcbf", "icon": "🟡"},
    "LOW":      {"color": "#276749", "bg": "#f0fff4", "icon": "🟢"},
}


def scan_keywords(text):
    """Return dict of found crime keywords with metadata."""
    text_lower = text.lower()
    found = {}
    for kw, (sev, cat) in CRIME_KEYWORDS.items():
        idx = text_lower.find(kw)
        if idx != -1:
            s = max(0, idx - 60)
            e = min(len(text), idx + len(kw) + 60)
            snippet = (
                ("…" if s > 0 else "")
                + text[s:e].replace("\n", " ").strip()
                + ("…" if e < len(text) else "")
            )
            found[kw] = {
                "severity": sev,
                "category": cat,
                "count":    text_lower.count(kw),
                "snippet":  snippet,
            }
    return found


def get_legal_measures(keywords_list):
    """Call the same API to get a lawyer-grade advisory JSON."""
    kw_str = ", ".join(keywords_list)
    prompt = f"""You are a senior criminal lawyer. These serious offence keywords were found 
    in a legal document: {kw_str}.
    
    Generate a structured legal advisory in JSON with exactly these keys:
    - "case_type": string (overall case classification)
    - "urgency": string — one of IMMEDIATE / HIGH / MODERATE
    - "applicable_laws": list of objects with "law" and "plain_meaning" (1 sentence each)
    - "immediate_actions": list of strings (max 6 — what lawyer/client must do NOW)
    - "evidence_checklist": list of strings (max 6 — what evidence to collect/preserve)
    - "bail_status": string (bailable / non-bailable / court discretion + 1-line reason)
    - "court_steps": list of strings (max 5 — steps from FIR/filing to trial)
    - "lawyer_note": string (3-sentence plain-English advice for the client)
    
    Respond ONLY in valid JSON. No preamble, no markdown, no extra text."""

    # ensure API key is available before making any API calls
    if not API_KEY:
        return {
            "error": "API key not configured. Set GITHUB_TOKEN in your local .env file "
                     "or in Streamlit secrets, then rerun the app."
        }

    headers = {
        "Content-Type":  "application/json",
        "Accept":        "application/json",
        "Authorization": f"Bearer {API_KEY}",
    }
    payload = {
        "model":       MODEL_NAME,
        "max_tokens":  1800,
        "temperature": 0.2,
        "messages":    [{"role": "user", "content": prompt}],
    }
    try:
        resp = requests.post(API_URL, headers=headers, json=payload)
        if resp.status_code == 200:
            raw   = resp.json()["choices"][0]["message"]["content"]
            clean = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
            return json.loads(clean)
        if resp.status_code == 401:
            return {
                "error": "Authentication failed with status code 401. "
                         "Please verify that your GITHUB_TOKEN is correct and has access "
                         "to GitHub Models / Azure Inference."
            }
        return {"error": f"API error {resp.status_code}"}
    except Exception as ex:
        return {"error": str(ex)}


def build_legal_measures_html(rpt, kws, bail_col):
    today    = datetime.now().strftime("%B %d, %Y")
    urgency  = rpt.get("urgency", "HIGH")
    bail_txt = rpt.get("bail_status", "To be determined.")
    kw_tags  = "".join(
        f"<span style='background:#1e3a5f;color:#f6e05e;border-radius:4px;"
        f"padding:2px 9px;font-size:.74rem;margin:2px;display:inline-block;'>{k.title()}</span>"
        for k in kws
    )
    laws_rows = "".join(
        f"<tr><td><strong>{l.get('law','')}</strong></td><td>{l.get('plain_meaning','')}</td></tr>"
        for l in rpt.get("applicable_laws", [])
    )
    ia_li = "".join(f"<li>{a}</li>" for a in rpt.get("immediate_actions", []))
    ev_li = "".join(f"<li>{e}</li>" for e in rpt.get("evidence_checklist", []))
    ct_li = "".join(f"<li><strong>{i+1}.</strong> {s}</li>" for i, s in enumerate(rpt.get("court_steps", [])))

    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>Legal Measures Report — {today}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=Source+Serif+4:wght@300;400;600&display=swap');
:root{{--n:#1e3a5f;--g:#d4af37;--cr:#fafaf7;}}
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:'Source Serif 4',Georgia,serif;background:var(--cr);color:#2a2a2a;font-size:14px;line-height:1.8;}}
.cover{{background:linear-gradient(150deg,var(--n),#152c47);color:white;padding:56px 65px 48px;}}
.cover h1{{font-family:'Playfair Display',serif;font-size:2.4rem;margin:12px 0 6px;}}
.gl{{width:55px;height:3px;background:var(--g);margin:16px 0 18px;}}
.lb{{font-size:.62rem;letter-spacing:.22em;text-transform:uppercase;color:var(--g);margin-bottom:3px;}}
.vl{{font-size:.92rem;font-weight:600;margin-bottom:12px;}}
.cf{{background:rgba(212,175,55,.12);border:1px solid rgba(212,175,55,.35);color:rgba(255,255,255,.8);
     font-size:.62rem;letter-spacing:.24em;text-transform:uppercase;text-align:center;padding:6px;margin-top:22px;}}
.body{{max-width:820px;margin:0 auto;padding:42px 56px;}}
.nc{{background:linear-gradient(135deg,var(--n),#2c5282);color:white;border-radius:10px;
     padding:20px 24px;margin:20px 0;border-left:5px solid var(--g);font-size:.96rem;line-height:1.85;}}
.bb{{border-radius:8px;padding:12px 18px;font-weight:700;font-size:.92rem;margin:12px 0;
     border:2px solid {bail_col};color:{bail_col};background:{bail_col}18;}}
.grid{{display:grid;grid-template-columns:1fr 1fr;gap:18px;margin-top:18px;}}
.sec{{background:white;border:1px solid #e2d9c8;border-radius:10px;padding:15px 19px;break-inside:avoid;margin-bottom:16px;}}
.sec h3{{font-family:'Playfair Display',serif;font-size:.98rem;color:var(--n);
         border-bottom:2px solid var(--g);padding-bottom:7px;margin-bottom:11px;}}
.sec ul{{margin:0;padding-left:16px;}}
.sec li{{padding:3px 0;border-bottom:1px solid #f0ebe0;font-size:.88rem;}}
.sec li:last-child{{border:none;}}
table{{width:100%;border-collapse:collapse;}}
th{{background:var(--n);color:white;padding:8px 13px;font-size:.78rem;letter-spacing:.08em;text-align:left;}}
td{{padding:8px 13px;border-bottom:1px solid #ede8dc;font-size:.87rem;vertical-align:top;}}
tr:nth-child(even) td{{background:#f9f6f0;}}
.disc{{background:#fffbeb;border:1px solid #f6e05e;border-radius:8px;
       padding:13px 17px;font-size:.8rem;color:#744210;margin-top:20px;}}
.ft{{background:var(--n);color:rgba(255,255,255,.5);text-align:center;
     padding:15px;font-size:.7rem;letter-spacing:.05em;margin-top:44px;}}
.ft strong{{color:var(--g);}}
</style></head><body>
<div class="cover">
  <div style="font-size:.62rem;letter-spacing:.22em;text-transform:uppercase;background:rgba(212,175,55,.18);
       border:1px solid var(--g);color:#f0d080;display:inline-block;padding:3px 11px;border-radius:2px;">
    Privileged Legal Document</div>
  <h1>Legal Measures Report</h1>
  <div class="gl"></div>
  <p style="opacity:.7;font-size:.95rem;margin-bottom:18px;">AI Criminal Case Advisory — {today}</p>
  <div class="lb">Detected Keywords</div><div style="margin-bottom:16px;">{kw_tags}</div>
  <div class="lb">Case Type</div><div class="vl">{rpt.get('case_type','—')}</div>
  <div class="lb">Urgency</div><div class="vl" style="color:#fc8181;">{urgency}</div>
  <div class="cf">Attorney–Client Privilege &nbsp;|&nbsp; Work Product Doctrine &nbsp;|&nbsp; Do Not Distribute</div>
</div>
<div class="body">
  <div class="nc">💼 <strong>Lawyer's Note</strong><br><br>{rpt.get('lawyer_note','')}</div>
  <div class="bb">🔒 Bail Status: {bail_txt}</div>
  <div class="sec">
    <h3>📜 Applicable Laws &amp; Sections</h3>
    <table><thead><tr><th>Law / Section</th><th>Plain Meaning</th></tr></thead>
    <tbody>{laws_rows or '<tr><td colspan="2">No laws identified.</td></tr>'}</tbody></table>
  </div>
  <div class="grid">
    <div class="sec"><h3>⚡ Immediate Actions</h3><ul>{ia_li or '<li>None</li>'}</ul></div>
    <div class="sec"><h3>🧾 Evidence Checklist</h3><ul>{ev_li or '<li>None</li>'}</ul></div>
  </div>
  <div class="sec"><h3>🏛️ Court Steps (Filing → Trial)</h3><ul>{ct_li or '<li>None</li>'}</ul></div>
  <div class="disc">⚠️ <strong>Disclaimer:</strong> AI-generated for informational purposes only.
    Not formal legal advice. Always consult a qualified attorney before taking any legal action.</div>
</div>
<div class="ft">
  <strong>Legal Document Assistant</strong> &nbsp;|&nbsp; {today} &nbsp;|&nbsp; <em>Confidential</em>
</div>
</body></html>"""


def build_law_firm_html(analysis, firm_name, attorney_name, client_name, matter_ref):
    """Build branded HTML report for law firm."""
    today        = datetime.now().strftime("%B %d, %Y")
    score        = min(100, len(analysis.get("risks", [])) * 15)
    score_color  = "#38a169" if score <= 30 else ("#d69e2e" if score <= 65 else "#c53030")
    score_label  = "LOW" if score <= 30 else ("MODERATE" if score <= 65 else "HIGH")

    risks_rows   = "".join(
        f"<tr><td>{i+1}</td><td>{r}</td><td class='bh'>High</td></tr>"
        for i, r in enumerate(analysis.get("risks", []))
    )
    oblig_html   = "".join(f"<li>{o}</li>" for o in analysis.get("obligations", []))
    clauses_html = "".join(
        f"<div class='cc'><div class='ct'>{c.get('title','')}</div>"
        f"<div class='cb'>{c.get('explanation','')}</div></div>" 
        for c in analysis.get("clauses", [])
    )
    terms_rows   = "".join(
        f"<tr><td><strong>{t.get('term','')}</strong></td><td>{t.get('definition','')}</td></tr>"
        for t in analysis.get("keyTerms", [])
    )

    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>Legal Report — {firm_name or 'Law Firm'}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=Source+Serif+4:wght@300;400;600&display=swap');
:root{{--n:#1a2d4a;--g:#c9a84c;--g2:#f0d080;--cr:#fafaf7;}}
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:'Source Serif 4',Georgia,serif;background:var(--cr);color:#2a2a2a;font-size:14px;line-height:1.75;}}
.cover{{background:linear-gradient(160deg,var(--n) 0%,#0f1e30 100%);color:white;padding:70px 80px 60px;}}
.cover h1{{font-family:'Playfair Display',serif;font-size:2.8rem;margin-bottom:8px;}}
.gl{{width:65px;height:3px;background:linear-gradient(90deg,var(--g),var(--g2));margin:20px 0;}}
.tag{{display:inline-block;background:rgba(201,168,76,.22);border:1px solid var(--g);color:var(--g2);
      font-size:.7rem;letter-spacing:.22em;text-transform:uppercase;padding:4px 14px;border-radius:2px;margin-bottom:24px;}}
.meta{{display:grid;grid-template-columns:1fr 1fr;gap:12px 40px;max-width:500px;
       border-top:1px solid rgba(201,168,76,.28);padding-top:20px;margin-top:28px;}}
.meta .lb{{font-size:.66rem;letter-spacing:.2em;text-transform:uppercase;color:var(--g);margin-bottom:2px;}}
.meta .vl{{font-size:.93rem;font-weight:600;}}
.cf{{background:rgba(201,168,76,.12);border:1px solid rgba(201,168,76,.35);color:rgba(255,255,255,.8);
     font-size:.68rem;letter-spacing:.24em;text-transform:uppercase;text-align:center;padding:7px;margin-top:30px;}}
.pg{{max-width:860px;margin:0 auto;padding:50px 70px;}}
.sec{{margin-bottom:44px;page-break-inside:avoid;}}
.sl{{font-size:.66rem;letter-spacing:.25em;text-transform:uppercase;color:var(--g);margin-bottom:5px;}}
.st{{font-family:'Playfair Display',serif;font-size:1.5rem;color:var(--n);
     border-bottom:2px solid var(--g);padding-bottom:9px;margin-bottom:20px;}}
.rb{{width:108px;height:108px;border-radius:50%;border:5px solid {score_color};display:flex;
     flex-direction:column;align-items:center;justify-content:center;flex-shrink:0;
     background:white;box-shadow:0 4px 22px rgba(0,0,0,.08);}}
.rb .rn{{font-size:1.9rem;font-weight:800;color:{score_color};line-height:1;}}
.rb .rm{{font-size:.68rem;color:#718096;}}
.rb .rl{{font-size:.63rem;letter-spacing:.14em;text-transform:uppercase;color:{score_color};margin-top:2px;font-weight:700;}}
.sc{{background:linear-gradient(135deg,#eef5ff,#e8f4ed);border-left:5px solid var(--n);
     border-radius:0 10px 10px 0;padding:20px 26px;font-size:1rem;line-height:1.85;color:var(--n);}}
.cc{{border:1px solid #d8d0c0;border-radius:8px;margin-bottom:13px;overflow:hidden;}}
.ct{{background:var(--n);color:white;padding:9px 17px;font-family:'Playfair Display',serif;font-size:.97rem;font-weight:600;}}
.cb{{padding:13px 17px;font-size:.93rem;}}
.ol{{background:#f0fff4;border-left:5px solid #38a169;border-radius:0 8px 8px 0;padding:16px 22px;list-style:none;}}
.ol li{{padding:5px 0 5px 20px;border-bottom:1px solid #c6f6d5;color:#22543d;position:relative;}}
.ol li::before{{content:'✓';position:absolute;left:0;color:#38a169;font-weight:800;}}
.ol li:last-child{{border:none;}}
.rt{{width:100%;border-collapse:collapse;border-radius:8px;overflow:hidden;box-shadow:0 2px 10px rgba(0,0,0,.06);}}
.rt thead tr{{background:var(--n);color:white;}}
.rt th{{padding:11px 15px;text-align:left;font-size:.83rem;letter-spacing:.07em;}}
.rt td{{padding:11px 15px;font-size:.91rem;border-bottom:1px solid #e8e0d0;}}
.rt tr:nth-child(even) td{{background:#faf7f2;}}
.bh{{background:#fff5f5;color:#c53030;border:1px solid #feb2b2;border-radius:20px;
     padding:2px 11px;font-size:.73rem;font-weight:700;}}
.tt{{width:100%;border-collapse:collapse;}}
.tt th{{background:#f4f1ea;color:var(--n);padding:9px 15px;text-align:left;
        font-size:.8rem;letter-spacing:.1em;text-transform:uppercase;border-bottom:2px solid var(--g);}}
.tt td{{padding:9px 15px;border-bottom:1px solid #d8d0c0;font-size:.91rem;vertical-align:top;}}
.disc{{background:#fffbeb;border:1px solid #f6e05e;border-radius:8px;padding:17px 20px;font-size:.8rem;color:#744210;margin-top:36px;}}
.ft{{background:var(--n);color:rgba(255,255,255,.6);text-align:center;padding:20px 40px;
     font-size:.76rem;letter-spacing:.05em;margin-top:55px;}}
.ft strong{{color:var(--g);}}
</style></head><body>
<div class="cover">
  <div class="tag">Privileged &amp; Confidential</div>
  <h1>Legal Document<br>Analysis Report</h1>
  <div class="gl"></div>
  <div class="meta">
    <div><div class="lb">Prepared By</div><div class="vl">{attorney_name or 'Legal Document Assistant'}</div></div>
    <div><div class="lb">Law Firm</div><div class="vl">{firm_name or '—'}</div></div>
    <div><div class="lb">Prepared For</div><div class="vl">{client_name or '—'}</div></div>
    <div><div class="lb">Matter Reference</div><div class="vl">{matter_ref or '—'}</div></div>
    <div><div class="lb">Date</div><div class="vl">{today}</div></div>
    <div><div class="lb">Risk Level</div><div class="vl" style="color:{score_color};">{score_label} RISK</div></div>
  </div>
  <div class="cf">Attorney–Client Privilege &nbsp;|&nbsp; Work Product Doctrine &nbsp;|&nbsp; Do Not Distribute</div>
</div>
<div class="pg">
  <div class="sec">
    <div class="sl">Overview</div><div class="st">Risk Assessment</div>
    <div style="display:flex;align-items:center;gap:26px;margin-bottom:24px;">
      <div class="rb"><div class="rn">{score}</div><div class="rm">/100</div><div class="rl">{score_label}</div></div>
      <div style="font-size:.97rem;color:#5a6470;line-height:1.7;">
            Evaluated across <strong>{len(analysis.get('risks',[]))} risk factors</strong>,
            <strong>{len(analysis.get('clauses',[]))} clauses</strong>, and
            <strong>{len(analysis.get('obligations',[]))} obligations</strong>.<br>
            Risk score: <strong style="color:{score_color};">{score}/100 ({score_label})</strong>
          </div>
    </div>
  </div>
  <div class="sec"><div class="sl">Section 1</div><div class="st">Executive Summary</div>
    <div class="sc">{analysis.get('summary','No summary available.')}</div></div>
  <div class="sec"><div class="sl">Section 2</div><div class="st">Key Clauses</div>
    {clauses_html or '<p>No clauses identified.</p>'}</div>
  <div class="sec"><div class="sl">Section 3</div><div class="st">Obligations</div>
    <ul class="ol">{oblig_html or '<li>No obligations identified.</li>'}</ul></div>
  <div class="sec"><div class="sl">Section 4</div><div class="st">Risk Register</div>
    <table class="rt"><thead><tr><th>#</th><th>Risk</th><th>Severity</th></tr></thead>
        <tbody>{risks_rows or '<tr><td colspan="3">No risks identified.</td></tr>'}</tbody></table></div>
  <div class="sec"><div class="sl">Section 5</div><div class="st">Legal Glossary</div>
    <table class="tt"><thead><tr><th>Term</th><th>Definition</th></tr></thead>
        <tbody>{terms_rows or '<tr><td colspan="2">No terms defined.</td></tr>'}</tbody></table></div>
  <div class="disc"><strong>Disclaimer:</strong> AI-generated for informational purposes only.
        Not formal legal advice. Consult a qualified attorney before taking any legal action.</div>
</div>
<div class="ft">
      {(('<strong>' + firm_name + '</strong> &nbsp;|&nbsp; ') if firm_name else '')}
      Legal Document Assistant &nbsp;|&nbsp; {today} &nbsp;|&nbsp; <em>Confidential</em>
    </div>
</body></html>"""

st.markdown("""
<style>

/* Google Font */
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');

*{
font-family:'Poppins',sans-serif;
}

/* MAIN BACKGROUND */
.main{
background: linear-gradient(135deg,#eef2ff 0%,#e0e7ff 50%,#f8fafc 100%);
}

/* HEADER */
.main-header{
font-size:3.5rem;
font-weight:700;
background: linear-gradient(135deg,#6366f1 0%,#8b5cf6 50%,#ec4899 100%);
-webkit-background-clip:text;
-webkit-text-fill-color:transparent;
text-align:center;
margin-bottom:0.5rem;
}

.sub-header{
font-size:1.3rem;
color:#475569;
text-align:center;
margin-bottom:2rem;
font-weight:300;
}

/* SECTION TITLES */

.section-header{
font-size:2rem;
font-weight:600;
background: linear-gradient(90deg,#4f46e5 0%,#7c3aed 100%);
-webkit-background-clip:text;
-webkit-text-fill-color:transparent;
border-bottom:3px solid #6366f1;
padding-bottom:0.5rem;
margin-top:2rem;
margin-bottom:1.5rem;
}

/* HERO CARD */

.hero-card{
background: linear-gradient(135deg,#4f46e5 0%,#7c3aed 50%,#9333ea 100%);
padding:1.5rem;
border-radius:15px;
color:white;
box-shadow:0 10px 30px rgba(79,70,229,0.35);
margin-bottom:2rem;
}

/* INFO BOX */

.info-box{
background:linear-gradient(135deg,#eff6ff 0%,#dbeafe 100%);
color:#1e3a8a;
padding:1.5rem;
border-radius:15px;
border-left:6px solid #2563eb;
margin:1rem 0;
box-shadow:0 6px 20px rgba(0,0,0,0.08);
}

/* SUCCESS BOX */

.success-box{
background:linear-gradient(135deg,#ecfdf5 0%,#bbf7d0 100%);
color:#064e3b;
padding:1.5rem;
border-radius:15px;
border-left:6px solid #059669;
margin:1rem 0;
}

/* RISK BOX */

.risk-box{
background:linear-gradient(135deg,#fef2f2 0%,#fecaca 100%);
color:#7f1d1d;
padding:1.5rem;
border-radius:15px;
border-left:6px solid #dc2626;
margin:1rem 0;
}

/* BUTTONS */

.stButton>button{
background:linear-gradient(135deg,#2563eb 0%,#7c3aed 100%);
color:white;
border:none;
border-radius:25px;
padding:0.75rem 2rem;
font-weight:600;
font-size:1.1rem;
box-shadow:0 6px 20px rgba(79,70,229,0.4);
transition:all 0.3s ease;
}

.stButton>button:hover{
transform:translateY(-3px);
box-shadow:0 10px 25px rgba(79,70,229,0.6);
}

/* EXPANDER */

.streamlit-expanderHeader{
background:linear-gradient(135deg,#f3f4f6 0%,#e5e7eb 100%);
border-radius:10px;
font-weight:600;
padding:1rem;
border-left:5px solid #6366f1;
}

/* FILE UPLOADER */

[data-testid="stFileUploader"]{
background:white;
padding:2rem;
border-radius:15px;
border:2px dashed #6366f1;
transition:all 0.3s ease;
}

[data-testid="stFileUploader"]:hover{
border-color:#8b5cf6;
background:#f9fafb;
}

/* TERM CARD */

.term-card{
background:white;
padding:1.5rem;
border-radius:12px;
box-shadow:0 4px 15px rgba(0,0,0,0.08);
margin:1rem 0;
border-top:4px solid #6366f1;
}

/* CHAT */

.chat-message{
padding:1rem;
border-radius:15px;
margin:1rem 0;
}

.user-message{
background:linear-gradient(135deg,#dbeafe 0%,#bfdbfe 100%);
border-left:5px solid #3b82f6;
margin-left:2rem;
}

.assistant-message{
background:linear-gradient(135deg,#f3e8ff 0%,#e9d5ff 100%);
border-left:5px solid #a855f7;
margin-right:2rem;
}

/* SIDEBAR */

[data-testid="stSidebar"]{
background:linear-gradient(180deg,#0f172a 0%,#1e293b 50%,#312e81 100%);
color:white;
}

/* FOOTER */

.footer{
background:linear-gradient(135deg,#0f172a 0%,#1e293b 50%,#312e81 100%);
padding:2rem;
border-radius:15px;
text-align:center;
color:white;
box-shadow:0 10px 30px rgba(0,0,0,0.3);
margin-top:3rem;
}

</style>
""", unsafe_allow_html=True)
# Header
st.markdown('<h1 class="main-header">⚖️ Legal Document Assistant</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">🔍 Understand complex legal documents in simple language</p>', 
unsafe_allow_html=True)

# Hero info box
st.markdown("""
<div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
            padding: 1.5rem; border-radius: 15px; color: white; 
            box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3); margin-bottom: 2rem;'>
    <h3 style='color: white; margin: 0 0 0.5rem 0;'>✨ How it works</h3>
    <p style='margin: 0; font-size: 1.1rem;'>
        Upload your legal document or paste text, and I'll break it down into simple terms. 
        Then ask me any questions you have!
    </p>
</div>
""", unsafe_allow_html=True)

# Sidebar
st.sidebar.markdown("""
<div style='text-align: center; padding: 1rem 0;'>
    <h1 style='font-size: 2.5rem; margin: 0;'>⚖️</h1>
    <h2 style='color: white; margin: 0.5rem 0;'>Instructions</h2>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div style='background: rgba(255,255,255,0.1); padding: 1rem; border-radius: 10px; margin: 1rem 0;'>
    <h3 style='color: #fbbf24;'>📋 How to use:</h3>
    <ol style='color: white; line-height: 2;'>
        <li><strong>Upload Document</strong><br/>Paste text or upload a file</li>
        <li><strong>Analyze</strong><br/>Click to get instant insights</li>
        <li><strong>Ask Questions</strong><br/>Interactive Q&A support</li>
    </ol>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div style='background: rgba(255,255,255,0.1); padding: 1rem; border-radius: 10px; margin: 1rem 0;'>
    <h3 style='color: #34d399;'>✨ Features:</h3>
    <ul style='color: white; line-height: 2;'>
        <li>📝 Simple summaries</li>
        <li>💬 Plain language explanations</li>
        <li>⚠️ Risk identification</li>
        <li>🎯 Deadline tracking</li>
        <li>💡 Interactive Q&A</li>
        <li>🔒 Privacy-focused</li>
    </ul>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div style='background: rgba(255,255,255,0.1); padding: 1rem; border-radius: 10px; margin: 1rem 0;'>
    <h3 style='color: #60a5fa;'>❓ Example Questions:</h3>
    <ul style='color: white; line-height: 1.8; font-size: 0.9rem;'>
        <li>What are my main responsibilities?</li>
        <li>Can I cancel this contract?</li>
        <li>What happens if I don't pay on time?</li>
        <li>What does [term] mean?</li>
    </ul>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style='background: rgba(245, 158, 11, 0.2); padding: 1rem; border-radius: 10px; border-left: 4px solid #f59e0b;'>
    <p style='color: white; margin: 0; font-size: 0.9rem;'>
        <strong>⚠️ Note:</strong> This tool provides information only. 
        For legal advice, consult a qualified lawyer.
    </p>
</div>
""", unsafe_allow_html=True)

# Session state
if 'document' not in st.session_state:
    st.session_state.document = ""
if 'analysis' not in st.session_state:
    st.session_state.analysis = None
if 'documents' not in st.session_state:
    st.session_state.documents = {}    # filename -> text
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = []
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = {}   # document name -> list of messages   
# initialize feature-related session keys
if 'kw_found' not in st.session_state:
    st.session_state.kw_found = {}
if 'kw_text' not in st.session_state:
    st.session_state.kw_text = ""
if 'lm_report' not in st.session_state:
    st.session_state.lm_report = None
if 'lm_kws' not in st.session_state:
    st.session_state.lm_kws = []
# Document input section
st.markdown('<h2 class="section-header">📄 Step 1: Enter Your Document</h2>', unsafe_allow_html=True)

tab1, tab2 = st.tabs(["📝 Paste Text", "📁 Upload File"])

with tab1:
    document_input = st.text_area(
        "Paste your legal document here:",
        height=300,
        placeholder="Copy and paste your contract, agreement, or legal document here...",
        key="doc_input"
    )
    if document_input:
        st.session_state.document = document_input

with tab2:
    uploaded_files = st.file_uploader(
        "Choose one or more files",
        type=['txt', 'doc', 'docx', 'pdf'],
        accept_multiple_files=True
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            try:
                document_text = ""
                filename = uploaded_file.name
                lower = filename.lower()
                if lower.endswith(".pdf"):
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            document_text += text + "\n"
                elif lower.endswith(".docx"):
                    doc = Document(uploaded_file)
                    document_text = "\n".join([p.text for p in doc.paragraphs])
                else:
                    document_text = uploaded_file.read().decode("utf-8", errors="ignore")

                if document_text.strip() == "":
                    st.error(f"❌ {filename} is empty or could not be processed.")
                else:
                    st.session_state.documents[filename] = document_text
                    st.success(f"✅ Loaded: {filename}")
                    with st.expander(f"Preview {filename}"):
                        st.text_area(
                            "Document preview",
                            document_text[:2000] + ("..." if len(document_text) > 2000 else ""),
                            height=150,
                            disabled=True,
                            label_visibility="collapsed",
                        )
            except Exception as e:
                st.error(f"❌ {filename} could not be read: {e}")
    else:
        st.info("💡 Upload one or more TXT, DOC, DOCX, or PDF files.")

st.markdown("---")

# Analyze button
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    if st.button("Analyze Document(s)", type="primary", use_container_width=True):
        documents_to_analyze = []
        # Collect pasted document
        if st.session_state.get("document"):
            documents_to_analyze.append({"name": "Pasted Document", "text": st.session_state.document})
        # Collect uploaded documents
        if st.session_state.get("documents"):
            for name, txt in st.session_state.documents.items():
                documents_to_analyze.append({"name": name, "text": txt})
        if documents_to_analyze:
            with st.spinner(f"Analyzing {len(documents_to_analyze)} document(s)... This may take 10-15 seconds per document..."):
                st.session_state.analysis_results = []
                for doc in documents_to_analyze:
                    text = doc["text"]
                    if len(text) > MAX_DOC_CHARS:
                        st.warning(
                            f"{doc['name']} is over {MAX_DOC_CHARS} characters; "
                            "it will be split into smaller pieces and processed in sequence." 
                        )
                    analysis_result = analyze_document(text)
                    # keep original text so we can power features like deadline tracking
                    st.session_state.analysis_results.append(
                        {"name": doc["name"], "result": analysis_result, "text": text}
                    )
            st.success("Analysis complete! Scroll down to see results.")
        else:
            st.warning("Please enter or upload a document first!")

# Analysis results (display outside of button handler)
if st.session_state.analysis_results:
    st.markdown("---")
    st.markdown('<h2 class="section-header"> Analysis Results</h2>', unsafe_allow_html=True)
    for item in st.session_state.analysis_results:
        name = item.get("name", "Document")
        analysis = item.get("result", {})
        st.markdown(f"###  Document: {name}")
        if analysis.get("error"):
            st.error(f"❌ Error analyzing this document: {analysis['error']}")
            continue
        # Summary
        st.markdown("###  Summary")
        st.markdown(f'<div class="info-box"><strong> Overview:</strong><br/>{analysis.get("summary", "No summary available")}</div>', unsafe_allow_html=True)
        # Risk Meter
        risk_score = compute_risk_score(analysis)
        st.markdown("###  Risk Meter")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.progress(risk_score/100)
            risk_class = "risk-low" if risk_score < 40 else ("risk-medium" if risk_score < 70 else "risk-high")
            risk_emoji = "" if risk_score < 40 else ("" if risk_score < 70 else "🔴")
            st.markdown(f'<p class="risk-level {risk_class}">{risk_emoji} Overall risk level: {risk_score}%</p>', unsafe_allow_html=True)
        # Key Dates (use original document text for date extraction)
        original_text = item.get("text", "")
        dates = extract_dates_with_context(original_text)
        st.markdown("### 📅 Key Dates & Deadlines")
        render_deadline_tracker(dates)
        # Key Clauses
        if analysis.get("clauses"):
            st.markdown("### 📜 Key Clauses Explained")
            for i, clause in enumerate(analysis.get("clauses", []), 1):
                with st.expander(f"📌 {clause.get('title', f'Clause {i}')}"):
                    st.write(clause.get('explanation', 'No explanation available'))
        # Obligations
        if analysis.get("obligations"):
            st.markdown("### ✅ Your Obligations (What You Must Do)")
            st.markdown('<div class="success-box">', unsafe_allow_html=True)
            for i, obligation in enumerate(analysis.get("obligations", []), 1):
                st.markdown(f"**{i}.** {obligation}")
            st.markdown('</div>', unsafe_allow_html=True)
        # Key Terms
        if analysis.get("keyTerms"):
            st.markdown("### 📖 Key Terms Defined")
            cols = st.columns(2)
            for i, term in enumerate(analysis.get("keyTerms", [])):
                with cols[i % 2]:
                    st.markdown(f"""
                    <div class="term-card">
                        <h4 style='color: #667eea; margin: 0 0 0.5rem 0;'>🔑 {term.get('term', 'Term')}</h4>
                        <p style='margin: 0; color: #4a5568; font-style: italic;'>{term.get('definition', 'No definition available')}</p>
                    </div>
                    """, unsafe_allow_html=True)

#  FEATURE 2 — LAW FIRM HTML REPORT
#  Generate branded report from the analysis output


if st.session_state.analysis_results:
    # use first document's analysis by default
    _analysis = st.session_state.analysis_results[0]["result"]
else:
    _analysis = None

st.markdown("---")
st.markdown('<div class="feat-header">📄 Feature 2 — Law Firm HTML Report</div>', unsafe_allow_html=True)
st.write("Create a branded HTML summary for your law firm based on the AI analysis.")

firm_name     = st.text_input("🏢 Law Firm Name:",   placeholder="e.g. Harrison & Cole LLP",        key="firm_name")
attorney_name = st.text_input("👤 Attorney Name:",   placeholder="e.g. Sarah J. Harrison, Esq.",    key="attorney")
client_name   = st.text_input("🧑‍💼 Client Name:",  placeholder="e.g. Acme Corporation",           key="client")
matter_ref    = st.text_input("📁 Matter Reference:",placeholder="e.g. Matter No. 2024-0042",       key="matter")

if st.button("📄 Generate & Download Report", type="primary", use_container_width=True, key="gen_report"):
    if _analysis:
        html_out = build_law_firm_html(
            _analysis,
            firm_name, attorney_name, client_name, matter_ref,
        )
        fname = f"Legal_Report_{datetime.now().strftime('%Y_%m_%d')}.html"
        st.download_button(
            label="⬇️ Download Report",
            data=html_out,
            file_name=fname,
            mime="text/html",
            use_container_width=True,
            key="dl_report",
        )
        st.success("✅ Report ready! Download above → open in Chrome → Ctrl+P → Save as PDF.")
    else:
        st.warning("Please run an analysis first to generate a report.")

#  FEATURE 3 — CRIMINAL KEYWORD SCANNER
#  Detects 33 serious crime keywords with severity badges & context


st.markdown("---")
st.markdown('<div class="feat-header">🔍 Feature 3 — Criminal Keyword Scanner</div>', unsafe_allow_html=True)
st.write("Scans any text for serious criminal case keywords — murder, kidnapping, smuggling, hit and run, and 29 more.")

scan_text = st.text_area(
    "Text to scan (auto-filled from your document if available):",
    value=st.session_state.document or st.session_state.kw_text,
    height=140,
    placeholder="Paste case notes, FIR, legal document, or any text here...",
    key="scan_input",
)

if st.button("🔎 Scan for Criminal Keywords", type="primary", key="scan_btn"):
    if scan_text.strip():
        found = scan_keywords(scan_text)
        st.session_state["kw_found"] = found
        st.session_state["kw_text"]  = scan_text

        if found:
            total = len(found)
            crits = sum(1 for v in found.values() if v["severity"] == "CRITICAL")
            highs = sum(1 for v in found.values() if v["severity"] == "HIGH")
            meds  = sum(1 for v in found.values() if v["severity"] == "MEDIUM")
            st.error(f"⚠️ **{total} keyword(s) detected** — {crits} Critical · {highs} High · {meds} Medium")

            by_cat = {}
            for kw, info in found.items():
                by_cat.setdefault(info["category"], []).append((kw, info))

            for cat_name, items in by_cat.items():
                st.markdown(f"**📂 {cat_name}**")
                for kw, info in items:
                    sty = SEV_STYLE[info["severity"]]
                    st.markdown(
                        f'<div class="kw-card" style="background:{sty["bg"]};border-left-color:{sty["color"]};">'
                        f'  <span style="font-weight:700;font-size:1rem;color:{sty["color"]};">'
                        f'    {sty["icon"]} {kw.title()}'
                        f'  </span>'
                        f'  <span class="sev-badge" style="background:{sty["color"]};">{info["severity"]}</span>'
                        f'  <span style="font-size:.8rem;color:#718096;float:right;">Found {info["count"]}×</span>'
                        f'  <div style="font-size:.8rem;color:#555;margin-top:2px;">Category: {info["category"]}</div>'
                        f'  <div class="kw-snip">"{info["snippet"]}"</div>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
        else:
            st.success("✅ No serious criminal case keywords found in this text.")
    else:
        st.warning("Please paste or upload text first.")


#  FEATURE 4 — AI LEGAL MEASURES REPORT
#  Takes detected keywords → same API → lawyer advisory + download


st.markdown("---")
st.markdown('<div class="feat-header">⚖️ Feature 4 — AI Legal Measures Report</div>', unsafe_allow_html=True)
st.write("Generates a lawyer-grade advisory from detected keywords — applicable laws, bail status, evidence checklist, court steps.")

def _bail_colour(txt):
    txt = txt.lower()
    if "non-bailable" in txt:
        return "#c53030"
    if "bailable" in txt:
        return "#276749"
    return "#b7791f"

default_kws = ", ".join(st.session_state.get("kw_found", {}).keys())
kw_input = st.text_input(
    "Keywords (auto-filled from scanner, or type your own):",
    value=default_kws,
    placeholder="e.g. murder, kidnapping, smuggling",
    key="kw_input",
)

if st.button("📋 Generate Legal Measures Report", type="primary", key="gen_measures"):
    kw_list = [k.strip() for k in kw_input.split(",") if k.strip()]
    if kw_list:
        with st.spinner("Consulting AI legal database — building report..."):
            rpt = get_legal_measures(kw_list)
        if "error" in rpt:
            st.error(f"Could not generate report: {rpt['error']}")
        else:
            st.session_state["lm_report"] = rpt
            st.session_state["lm_kws"]    = kw_list
    else:
        st.warning("Please enter at least one keyword.")

if st.session_state.get("lm_report"):
    rpt      = st.session_state["lm_report"]
    kws      = st.session_state.get("lm_kws", [])
    urgency  = rpt.get("urgency", "HIGH")
    bail_txt = rpt.get("bail_status", "To be determined.")
    bail_col = _bail_colour(bail_txt)

    st.markdown("---")

    h1, h2, h3 = st.columns(3)
    with h1:
        st.markdown(f"**Case Type**  \n`{rpt.get('case_type','—')}`")
    with h2:
        urg_col = {"IMMEDIATE": "#c53030", "HIGH": "#c05621", "MODERATE": "#b7791f"}.get(urgency.upper(), "#b7791f")
        st.markdown(f"**Urgency**  \n<span style='color:{urg_col};font-weight:800;font-size:1.1rem;'>{urgency}</span>", unsafe_allow_html=True)
    with h3:
        st.markdown(f"**Bail Status**  \n<span style='color:{bail_col};font-weight:700;'>{bail_txt}</span>", unsafe_allow_html=True)

    st.markdown(
        f"<div style='background:linear-gradient(135deg,#1e3a5f,#2c5282);color:white;"
        f"border-radius:10px;padding:18px 22px;margin:14px 0;border-left:5px solid #d4af37;"
        f"font-size:.97rem;line-height:1.8;'>"
        f"💼 <strong>Lawyer's Note</strong><br><br>{rpt.get('lawyer_note','')}</div>",
        unsafe_allow_html=True,
    )

    left, right = st.columns(2)
    with left:
        st.markdown("**📜 Applicable Laws**")
        for law in rpt.get("applicable_laws", []):
            st.markdown(
                f"<div style='background:#f7f3ea;border-left:4px solid #d4af37;"
                f"border-radius:6px;padding:8px 12px;margin-bottom:6px;'>"
                f"<strong>{law.get('law','')}</strong><br>"
                f"<span style='font-size:.85rem;color:#4a5568;'>{law.get('plain_meaning','')}</span>"
                f"</div>",
                unsafe_allow_html=True,
            )
        st.markdown("**⚡ Immediate Actions**")
        for ia in rpt.get("immediate_actions", []):
            st.markdown(f"- {ia}")

    with right:
        st.markdown("**🧾 Evidence Checklist**")
        for ev in rpt.get("evidence_checklist", []):
            st.markdown(f"- ✅ {ev}")
        st.markdown("**🏛️ Court Steps**")
        for i, step in enumerate(rpt.get("court_steps", []), 1):
            st.markdown(f"**{i}.** {step}")

    # Download button
    html_measures = build_legal_measures_html(rpt, kws, bail_col)
    fname_m = f"Legal_Measures_{datetime.now().strftime('%Y_%m_%d')}.html"

    d1, d2, d3 = st.columns([1, 2, 1])
    with d2:
        st.download_button(
            label="⬇️ Download Legal Measures Report",
            data=html_measures,
            file_name=fname_m,
            mime="text/html",
            use_container_width=True,
            key="dl_measures",
        )
    st.success("✅ Report ready! Open in Chrome → Ctrl+P → Save as PDF.")


# Q&A Section
# determine available documents for chat
available_docs = []
if st.session_state.get("document"):
    available_docs.append("Pasted Document")
available_docs.extend(list(st.session_state.get("documents", {}).keys()))

if available_docs:
    st.markdown("---")
    st.markdown('<h2 class="section-header">💬 Ask Questions About Your Document</h2>', unsafe_allow_html=True)

    selected_doc = st.selectbox("Choose document to chat about", available_docs)
    # initialize history for this doc if missing
    if selected_doc not in st.session_state.chat_history:
        st.session_state.chat_history[selected_doc] = []
    history = st.session_state.chat_history[selected_doc]

    # Display chat history
    if history:
        st.markdown("### 💭 Conversation History:")
        for message in history:
            if message["role"] == "user":
                st.markdown(f"""
                <div class="chat-message user-message">
                    <strong>👤 You:</strong><br/>
                    {message['content']}
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="chat-message assistant-message">
                    <strong>🤖 Assistant:</strong><br/>
                    {message['content']}
                </div>
                """, unsafe_allow_html=True)
     # Chat input
    question = st.text_input(
        "💬 Type your question here:",
        placeholder="Example: What are my main responsibilities in this contract?",
        key="question_input"
    )

    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("📨 Send Question", type="primary"):
            if question:
                with st.spinner("🤔 Thinking..."):
                    # choose document text
                    if selected_doc == "Pasted Document":
                        doc_text = st.session_state.document
                    else:
                        doc_text = st.session_state.documents.get(selected_doc, "")
                    if len(doc_text) > MAX_DOC_CHARS:
                        st.warning(
                            "Document is too long for the chat context; it will be "
                            "truncated before sending."
                        )
                    answer = ask_question(
                        doc_text,
                        question,
                        history
                    )
                    history.append({"role": "user", "content": question})
                    history.append({"role": "assistant", "content": answer})
                    st.session_state.chat_history[selected_doc] = history
                    st.rerun()
            else:
                st.warning("⚠️ Please type a question first!")

    with col2:
        if st.button("🗑️ Clear Chat"):
            st.session_state.chat_history[selected_doc] = []
            st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style='background: linear-gradient(135deg, #1e293b 0%, #334155 100%); 
            padding: 2rem; border-radius: 15px; text-align: center; 
            box-shadow: 0 10px 30px rgba(0,0,0,0.2); margin-top: 3rem;'>
    <h3 style='color: white; margin: 0 0 1rem 0;'>🔒 Privacy First</h3>
    <p style='color: #cbd5e1; margin: 0.5rem 0;'>
        Your documents are processed securely and <strong>not stored</strong> on our servers.
    </p>
    <p style='color: #cbd5e1; margin: 0.5rem 0;'>
        This tool provides informational guidance only. For legal advice, consult a qualified attorney.
    </p>
</div>
""", unsafe_allow_html=True)

